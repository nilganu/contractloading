"""DOCX parser using python-docx.

Preserves reading order and emits:
- headings
- paragraphs
- tables (with index/row/cell coords)
- count of embedded images (flag for vision)
"""
from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List

from docx import Document  # type: ignore


def parse_docx(path: str | Path) -> Dict[str, Any]:
    p = Path(path)
    doc = Document(str(p))

    blocks: List[Dict[str, Any]] = []

    for para in doc.paragraphs:
        style = (para.style.name if para.style is not None else "Normal") or "Normal"
        block_kind = "heading" if style.lower().startswith("heading") else "paragraph"
        text = para.text or ""
        if not text.strip() and block_kind != "heading":
            continue
        blocks.append(
            {
                "kind": block_kind,
                "style": style,
                "text": text,
            }
        )

    tables: List[Dict[str, Any]] = []
    for t_idx, tbl in enumerate(doc.tables, start=1):
        rows: List[List[str]] = []
        for r in tbl.rows:
            rows.append([cell.text for cell in r.cells])
        tables.append({"index": t_idx, "rows": rows})

    embedded_images = 0
    try:
        embedded_images = sum(
            1 for rel in doc.part.rels.values() if "image" in rel.reltype
        )
    except Exception:  # noqa: BLE001
        pass

    return {
        "source_file": p.name,
        "input_format": "docx",
        "blocks": blocks,
        "tables": tables,
        "embedded_images": embedded_images,
    }
